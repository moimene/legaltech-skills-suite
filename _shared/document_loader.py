"""
Document Loader - Utilidad compartida para carga segura de documentos.

Soporta: PDF, DOCX, DOC, TXT, EML, MSG
Ejecuta en TEE sin acceso a red.
"""

from pathlib import Path
from typing import Optional
from dataclasses import dataclass
from datetime import datetime
import re


@dataclass
class DocumentMetadata:
    """Metadatos extraídos del documento."""
    filename: str
    extension: str
    size_bytes: int
    created: Optional[datetime] = None
    modified: Optional[datetime] = None
    author: Optional[str] = None
    title: Optional[str] = None
    pages: Optional[int] = None


@dataclass
class LoadedDocument:
    """Documento cargado con texto y metadatos."""
    text: str
    metadata: DocumentMetadata
    raw_bytes: Optional[bytes] = None


class DocumentLoader:
    """Cargador seguro de documentos en Enclave TEE."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.eml', '.msg'}
    
    @classmethod
    def load(cls, file_path: str) -> LoadedDocument:
        """
        Carga documento de forma segura.
        
        Args:
            file_path: Ruta al documento
            
        Returns:
            LoadedDocument con texto y metadatos
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Extensión no soportada: {ext}")
        
        # Metadatos base
        stat = path.stat()
        metadata = DocumentMetadata(
            filename=path.name,
            extension=ext,
            size_bytes=stat.st_size,
            modified=datetime.fromtimestamp(stat.st_mtime)
        )
        
        # Cargar según tipo
        if ext == '.txt':
            text = cls._load_txt(path)
        elif ext == '.pdf':
            text, metadata = cls._load_pdf(path, metadata)
        elif ext in ['.docx', '.doc']:
            text, metadata = cls._load_docx(path, metadata)
        elif ext in ['.eml', '.msg']:
            text = cls._load_email(path)
        else:
            raise ValueError(f"Loader no implementado para: {ext}")
        
        return LoadedDocument(text=text, metadata=metadata)
    
    @staticmethod
    def _load_txt(path: Path) -> str:
        """Carga archivo de texto plano."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for enc in encodings:
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise ValueError("No se pudo decodificar el archivo de texto")
    
    @staticmethod
    def _load_pdf(path: Path, metadata: DocumentMetadata) -> tuple[str, DocumentMetadata]:
        """Carga PDF con PyMuPDF."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise RuntimeError("PyMuPDF no instalado: pip install pymupdf")
        
        doc = fitz.open(str(path))
        
        # Extraer metadatos
        meta = doc.metadata
        if meta:
            metadata.author = meta.get('author')
            metadata.title = meta.get('title')
            metadata.created = meta.get('creationDate')
        metadata.pages = len(doc)
        
        # Extraer texto
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        
        doc.close()
        return '\n'.join(text_parts), metadata
    
    @staticmethod
    def _load_docx(path: Path, metadata: DocumentMetadata) -> tuple[str, DocumentMetadata]:
        """Carga DOCX/DOC con python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx no instalado: pip install python-docx")
        
        doc = Document(str(path))
        
        # Extraer metadatos del core properties
        core = doc.core_properties
        if core:
            metadata.author = core.author
            metadata.title = core.title
            metadata.created = core.created
            metadata.modified = core.modified
        
        # Extraer texto de párrafos
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return '\n'.join(paragraphs), metadata
    
    @staticmethod
    def _load_email(path: Path) -> str:
        """Carga email EML/MSG."""
        import email
        from email import policy
        
        with open(path, 'rb') as f:
            msg = email.message_from_binary_file(f, policy=policy.default)
        
        parts = []
        parts.append(f"From: {msg.get('From', 'N/A')}")
        parts.append(f"To: {msg.get('To', 'N/A')}")
        parts.append(f"Date: {msg.get('Date', 'N/A')}")
        parts.append(f"Subject: {msg.get('Subject', 'N/A')}")
        parts.append("---")
        
        # Extraer cuerpo
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == 'text/plain':
                    parts.append(part.get_content())
        else:
            parts.append(msg.get_content())
        
        return '\n'.join(parts)


class LegalSegmenter:
    """Segmentador inteligente para documentos legales."""
    
    # Patrones de secciones legales
    PATTERNS = {
        'hechos': r'(?:HECHO\s+)?(?:PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|SÉPTIMO|OCTAVO|NOVENO|DÉCIMO)',
        'fundamentos': r'(?:FUNDAMENTO\s+)?(?:DE\s+DERECHO\s+)?(?:PRIMERO|SEGUNDO|TERCERO)',
        'romano': r'^[IVXLCDM]+\.?\-?\s',
        'numerico': r'^\d+\.?\-?\s',
        'letra': r'^[A-Z]\)\s',
    }
    
    @classmethod
    def segment(cls, text: str) -> list[tuple[int, str, str]]:
        """
        Segmenta texto legal en párrafos lógicos.
        
        Returns:
            Lista de (segment_id, tipo, texto)
        """
        combined = '|'.join(f'({p})' for p in cls.PATTERNS.values())
        
        segments = []
        current = []
        current_type = 'INTRO'
        segment_id = 0
        
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Detectar tipo de sección
            for name, pattern in cls.PATTERNS.items():
                if re.match(pattern, line, re.IGNORECASE):
                    # Guardar segmento anterior
                    if current:
                        segments.append((segment_id, current_type, ' '.join(current)))
                        segment_id += 1
                    current = [line]
                    current_type = name.upper()
                    break
            else:
                current.append(line)
        
        # Último segmento
        if current:
            segments.append((segment_id, current_type, ' '.join(current)))
        
        return segments
